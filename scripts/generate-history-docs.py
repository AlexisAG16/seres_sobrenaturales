from html.parser import HTMLParser
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT.parent / "VampinfosHTML"
OUTPUT = ROOT / "storage" / "history-docs"
IMAGES = SOURCE / "imagenes"

BURGUNDY = RGBColor(128, 18, 43)
VIOLET = RGBColor(91, 55, 120)
INK = RGBColor(38, 31, 39)
MUTED = RGBColor(105, 95, 104)
PARCHMENT = "F5F0E8"


class StoryParser(HTMLParser):
    block_tags = {"h2", "h3", "h4", "p", "li", "cite", "figcapture"}

    def __init__(self):
        super().__init__()
        self.blocks = []
        self.current = None
        self.buffer = []

    def handle_starttag(self, tag, attrs):
        if tag in self.block_tags:
            self.current = tag
            self.buffer = []
        elif tag == "br" and self.current:
            self.buffer.append("\n")

    def handle_data(self, data):
        if self.current:
            self.buffer.append(data)

    def handle_endtag(self, tag):
        if self.current == tag:
            text = " ".join("".join(self.buffer).split())
            if text:
                self.blocks.append((tag, text))
            self.current = None
            self.buffer = []


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 30, BURGUNDY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 20, BURGUNDY, 18, 8),
        ("Heading 2", 15, VIOLET, 14, 6),
        ("Heading 3", 12, BURGUNDY, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Story Quote" not in styles:
        quote = styles.add_style("Story Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["Story Quote"]
    quote.font.name = "Georgia"
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = VIOLET
    quote.paragraph_format.left_indent = Inches(0.4)
    quote.paragraph_format.right_indent = Inches(0.4)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(10)

    header = section.header.paragraphs[0]
    header.text = "ARCHIVO SOBRENATURAL  |  CRÓNICAS"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.add_run("Archivo personal de Vampire   •   ")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED
    add_page_number(footer)
    return doc


def add_cover(doc, kicker, title, subtitle, image=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    run = p.add_run(kicker.upper())
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = VIOLET

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title)

    subtitle_p = doc.add_paragraph(style="Subtitle")
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run(subtitle)

    if image and image.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(image), width=Inches(3.7))

    band = doc.add_table(rows=1, cols=1)
    band.autofit = False
    band.columns[0].width = Inches(6.5)
    cell = band.cell(0, 0)
    set_cell_shading(cell, "2A202D")
    cell_p = cell.paragraphs[0]
    cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_p.paragraph_format.space_before = Pt(8)
    cell_p.paragraph_format.space_after = Pt(8)
    band_run = cell_p.add_run("EDICIÓN DEL ARCHIVO SOBRENATURAL")
    band_run.font.name = "Aptos"
    band_run.font.size = Pt(9)
    band_run.font.bold = True
    band_run.font.color.rgb = RGBColor(245, 240, 232)
    doc.add_page_break()


def add_image(doc, image_path, caption, width=5.8):
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED


def add_intro_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PARCHMENT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = INK


def build_vampires():
    doc = configure_document()
    add_cover(doc, "Crónica de linajes", "Historia de los Vampiros", "Geografía, origen y genealogía del linaje vampírico", IMAGES / "geografiavmp.png")
    add_intro_box(doc, "Los vampiros del Archivo Sobrenatural nacen de una tradición europea atravesada por la sangre, la inmortalidad y los secretos familiares.")
    doc.add_heading("Territorios de origen", level=1)
    doc.add_paragraph("El origen del linaje se sitúa principalmente en Europa central y oriental. Rumania ocupa el centro de esta tradición, acompañada por Bulgaria, Hungría, Polonia, República Checa, Rusia y Ucrania.")
    add_image(doc, IMAGES / "geografiavmp.png", "Geografía principal de los orígenes vampíricos.", 5.9)
    doc.add_heading("Países vinculados al linaje", level=2)
    for country in ["Bulgaria", "Hungría", "Polonia", "República Checa", "Rumania — territorio principal", "Rusia", "Ucrania"]:
        doc.add_paragraph(country, style="List Bullet")
    doc.add_heading("Árbol genealógico", level=1)
    doc.add_paragraph("La genealogía se presenta en dos etapas, mostrando la expansión del linaje y sus conexiones a través del tiempo.")
    add_image(doc, IMAGES / "arbolv.png", "Primera etapa del árbol genealógico vampírico.", 6.25)
    add_image(doc, IMAGES / "arbolv2.png", "Segunda etapa del árbol genealógico vampírico.", 6.25)
    save(doc, "historia-vampiros.docx")


def build_brujas():
    doc = configure_document()
    add_cover(doc, "Crónica arcana", "Historia de las Brujas", "Origen europeo, persecución y legado de la magia", IMAGES / "origenesb.png")
    add_intro_box(doc, "La historia de las brujas está marcada menos por un único lugar de nacimiento que por siglos de relatos, persecuciones, conocimiento oculto y transmisión familiar.")
    doc.add_heading("Territorios y relatos de origen", level=1)
    doc.add_paragraph("Los relatos se extienden por gran parte de Europa central y occidental. En muchas regiones se conserva más memoria de la cacería y la persecución que de los primeros orígenes de la magia.")
    add_image(doc, IMAGES / "origenesb.png", "Mapa de los territorios asociados a los relatos de brujería.", 6.15)
    doc.add_heading("El conocimiento arcano", level=1)
    doc.add_paragraph("Las brujas del Archivo Sobrenatural practican magia blanca y oscura, elaboran pociones y transmiten sus capacidades mediante linajes. Algunas prolongan su vida y otras desarrollan poderes que superan los conocimientos de generaciones anteriores.")
    doc.add_heading("Árbol genealógico", level=1)
    doc.add_paragraph("El árbol conserva los vínculos familiares que permiten comprender cómo la magia pasa de una generación a otra y cómo nacen nuevas órdenes.")
    add_image(doc, IMAGES / "arbolb.png", "Árbol genealógico de las brujas.", 6.25)
    save(doc, "historia-brujas.docx")


def extract_necromancy_blocks():
    html = (SOURCE / "necroarcano.html").read_text(encoding="utf-8")
    parser = StoryParser()
    parser.feed(html)
    return parser.blocks


def build_necromancy():
    doc = configure_document()
    add_cover(doc, "Crónica mayor", "Necromancia vs. Arcano", "La historia de Mrs. Vampira y Manuela, la Bruja Suprema", IMAGES / "origenes.jpg")
    add_intro_box(doc, "Una historia de amistad, herencia, secretos familiares y venganza que comienza en el siglo XIX y enfrenta para siempre a la necromancia con el poder arcano.")
    add_image(doc, IMAGES / "origenes.jpg", "Primera representación de Mrs. Vampira.", 3.0)
    add_image(doc, IMAGES / "manuelareal.png", "Primera representación de Manuela, la Bruja Suprema.", 3.0)

    for tag, text in extract_necromancy_blocks():
        if text.lower() in {"origenes y leyendas necromanticas", "necromancia vs arcano , la historia"}:
            doc.add_heading(text.replace(" ,", ","), level=1)
        elif tag == "h2":
            doc.add_heading(text, level=1)
        elif tag == "h3":
            doc.add_heading(text, level=2)
        elif tag == "h4":
            doc.add_heading(text, level=2)
        elif tag == "li":
            doc.add_paragraph(text, style="List Bullet")
        elif tag == "cite":
            doc.add_paragraph(text, style="Story Quote")
        elif text != "FIN":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            ending = doc.add_paragraph()
            ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = ending.add_run("FIN")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = BURGUNDY
    save(doc, "necromancia-vs-arcano.docx")


def save(doc, filename):
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = filename.removesuffix(".docx").replace("-", " ").title()
    doc.core_properties.author = "Vampire — Archivo Sobrenatural"
    doc.core_properties.subject = "Historia del universo sobrenatural"
    doc.save(OUTPUT / filename)


if __name__ == "__main__":
    build_vampires()
    build_brujas()
    build_necromancy()
    print(f"Documents generated in {OUTPUT}")
