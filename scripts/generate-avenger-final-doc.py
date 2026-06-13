from pathlib import Path
from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT.parent / "EL FINAL DE AVENGER VAMPIRA.pdf"
OUTPUT = ROOT / "storage" / "character-docs" / "el-final-de-avenger-vampira.docx"

BURGUNDY = RGBColor(151, 22, 48)
VIOLET = RGBColor(104, 65, 129)
INK = RGBColor(38, 31, 39)
MUTED = RGBColor(105, 95, 104)


def set_font(run, name, size, color, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


reader = PdfReader(SOURCE)
text = "\n".join(page.extract_text() or "" for page in reader.pages)
lines = [line.strip() for line in text.splitlines() if line.strip()]
title = lines.pop(0)
paragraphs = []
current = ""
for line in lines:
    if current and (current.endswith((".", ":", "?", "!", "”")) or len(current) + len(line) > 650):
        paragraphs.append(current)
        current = line
    else:
        current = f"{current} {line}".strip()
if current:
    paragraphs.append(current)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.33

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(header.add_run("ARCHIVO SOBRENATURAL  |  HISTORIAS DE AVENGER"), "Aptos", 8, MUTED, True)
footer = section.footer.paragraphs[0]
set_font(footer.add_run("Vampire  •  Archivo Sobrenatural   "), "Aptos", 8, MUTED)
page_number(footer)

spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(44)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(10)
set_font(kicker.add_run("CRÓNICA FINAL"), "Aptos", 9, VIOLET, True)

heading = doc.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading.paragraph_format.space_after = Pt(8)
set_font(heading.add_run(title.title()), "Georgia", 27, BURGUNDY, True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(24)
set_font(subtitle.add_run("Las dos versiones sobre el destino de la heredera de Mrs. Vampira"), "Georgia", 12, MUTED, italic=True)

rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(24)
rule_pr = rule._p.get_or_add_pPr()
borders = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "18")
bottom.set(qn("w:color"), "971630")
borders.append(bottom)
rule_pr.append(borders)

for index, paragraph_text in enumerate(paragraphs):
    if paragraph_text.lower().startswith("hay dos versiones"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        set_font(p.add_run(paragraph_text), "Georgia", 13, VIOLET, True)
        continue
    if paragraph_text.lower().startswith("la primera"):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(6)
        set_font(h.add_run("La primera versión"), "Georgia", 16, BURGUNDY, True)
    elif paragraph_text.lower().startswith("la segunda"):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        set_font(h.add_run("La versión verdadera"), "Georgia", 16, BURGUNDY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.widow_control = True
    set_font(p.add_run(paragraph_text), "Calibri", 11, INK)

ending = doc.add_paragraph()
ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
ending.paragraph_format.space_before = Pt(22)
ending.paragraph_format.space_after = Pt(6)
set_font(ending.add_run("FIN"), "Georgia", 16, BURGUNDY, True)

doc.core_properties.title = "El final de Avenger Vampira"
doc.core_properties.author = "Vampire - Archivo Sobrenatural"
doc.core_properties.subject = "Historia de Avenger Vampira"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
